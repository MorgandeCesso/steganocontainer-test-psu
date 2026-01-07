#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Прототип стеганографии для DOCX:
- Шифрует полезную нагрузку (payload) с помощью AES‑GCM
  (пароль преобразуется в ключ через scrypt);
- Встраивает шифртекст в PNG‑картинку через LSB по RGB‑каналам;
- Помещает эту PNG‑картинку в DOCX как встроенное изображение (word/media/*).

Команды CLI:
- embed   — создать DOCX с изображением-стегаконтейнером;
- extract — достать скрытые данные из DOCX и расшифровать их.

Важно:
- Используется только PNG. JPEG для LSB категорически не подходит;
"""

import argparse
import os
import struct
import sys
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Optional

from PIL import Image
from docx import Document

from cryptography.hazmat.primitives.kdf.scrypt import Scrypt
from cryptography.hazmat.primitives.ciphers.aead import AESGCM


MAGIC = b"STEGDOC1"  # 8 байт «магической» сигнатуры
# Бинарный заголовок фиксированной длины:
# MAGIC(8) | SALT(16) | NONCE(12) | CT_LEN(4, uint32 LE) | CT(bytes)
HDR_FMT = "<8s16s12sI"
HDR_SIZE = struct.calcsize(HDR_FMT)


@dataclass
class CryptoPacket:
    salt: bytes
    nonce: bytes
    ciphertext: bytes

    def pack(self) -> bytes:
        return struct.pack(HDR_FMT, MAGIC, self.salt, self.nonce, len(self.ciphertext)) + self.ciphertext

    @staticmethod
    def unpack(blob: bytes) -> "CryptoPacket":
        if len(blob) < HDR_SIZE:
            raise ValueError("Blob too small for header")
        magic, salt, nonce, ct_len = struct.unpack(HDR_FMT, blob[:HDR_SIZE])
        if magic != MAGIC:
            raise ValueError("Magic mismatch: not a valid payload")
        if len(blob) < HDR_SIZE + ct_len:
            raise ValueError("Blob truncated: ciphertext length mismatch")
        ciphertext = blob[HDR_SIZE:HDR_SIZE + ct_len]
        return CryptoPacket(salt=salt, nonce=nonce, ciphertext=ciphertext)


def derive_key(password: str, salt: bytes, length: int = 32) -> bytes:
    kdf = Scrypt(
        salt=salt,
        length=length,
        n=2**15,   # CPU/mem cost
        r=8,
        p=1,
    )
    return kdf.derive(password.encode("utf-8"))


def encrypt(plaintext: bytes, password: str) -> CryptoPacket:
    salt = os.urandom(16)
    nonce = os.urandom(12)
    key = derive_key(password, salt)
    aesgcm = AESGCM(key)
    ciphertext = aesgcm.encrypt(nonce, plaintext, None)
    return CryptoPacket(salt=salt, nonce=nonce, ciphertext=ciphertext)


def decrypt(packet: CryptoPacket, password: str) -> bytes:
    key = derive_key(password, packet.salt)
    aesgcm = AESGCM(key)
    return aesgcm.decrypt(packet.nonce, packet.ciphertext, None)


def bytes_to_bits(data: bytes):
    for b in data:
        for i in range(7, -1, -1):
            yield (b >> i) & 1


def bits_to_bytes(bits):
    out = bytearray()
    cur = 0
    n = 0
    for bit in bits:
        cur = (cur << 1) | (bit & 1)
        n += 1
        if n == 8:
            out.append(cur)
            cur = 0
            n = 0
    if n != 0:
        # отбрасываем хвост неполных байт
        pass
    return bytes(out)


def lsb_capacity_bytes(img: Image.Image) -> int:
    """
    Возвращает вместимость в байтах при встраивании 1 бита на цветовой канал RGB.
    Используем только RGB (3 канала), по 1 младшему биту => 3 бита на пиксель.
    """
    w, h = img.size
    total_bits = w * h * 3
    return total_bits // 8


def embed_lsb_png(cover_png_path: str, payload: bytes, out_png_path: str) -> None:
    img = Image.open(cover_png_path)

    # Принудительно приводим к RGB (альфа‑канал игнорируем для простоты)
    img = img.convert("RGB")
    w, h = img.size
    pixels = list(img.getdata())

    cap = lsb_capacity_bytes(img)
    # Встраиваем длину полезной нагрузки (4 байта, uint32 LE) + сами данные
    blob = struct.pack("<I", len(payload)) + payload

    if len(blob) > cap:
        raise ValueError(f"Payload too large for cover image. Need {len(blob)} bytes, capacity {cap} bytes.")

    bit_iter = bytes_to_bits(blob)

    new_pixels = []
    try:
        for (r, g, b) in pixels:
            r = (r & ~1) | next(bit_iter)
            g = (g & ~1) | next(bit_iter)
            b = (b & ~1) | next(bit_iter)
            new_pixels.append((r, g, b))
    except StopIteration:
        # биты закончились — оставшиеся пиксели не трогаем
        new_pixels.extend(pixels[len(new_pixels):])

    stego = Image.new("RGB", (w, h))
    stego.putdata(new_pixels)
    stego.save(out_png_path, format="PNG")


def extract_lsb_png(stego_png_path: str) -> bytes:
    img = Image.open(stego_png_path).convert("RGB")
    pixels = list(img.getdata())

    bits = []
    for (r, g, b) in pixels:
        bits.append(r & 1)
        bits.append(g & 1)
        bits.append(b & 1)

    data = bits_to_bytes(bits)
    if len(data) < 4:
        raise ValueError("Not enough data in image")

    (payload_len,) = struct.unpack("<I", data[:4])
    payload = data[4:4 + payload_len]
    if len(payload) != payload_len:
        raise ValueError("Truncated payload in image (wrong length or not a stego image)")
    return payload


def create_docx_with_image(image_path: str, out_docx_path: str, title: Optional[str] = None) -> None:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    doc.add_paragraph("Изображение ниже содержит стегоконтейнер с данными. Для извлечения данных необходимо использовать специальную команду")
    doc.add_picture(image_path)  # default size
    doc.save(out_docx_path)


def extract_first_png_from_docx(docx_path: str) -> bytes:
    """
    Возвращает байты первого PNG‑файла, найденного в word/media/ внутри DOCX.
    """
    with zipfile.ZipFile(docx_path, "r") as zf:
        media_files = [n for n in zf.namelist() if n.startswith("word/media/") and n.lower().endswith(".png")]
        if not media_files:
            raise ValueError("No PNG images found in DOCX (word/media/*.png)")
        media_files.sort()
        return zf.read(media_files[0])


def cmd_embed(args: argparse.Namespace) -> None:
    plaintext = args.message.encode("utf-8") if args.message is not None else open(args.infile, "rb").read()
    packet = encrypt(plaintext, args.password)
    payload = packet.pack()

    embed_lsb_png(args.cover_png, payload, args.out_png)

    create_docx_with_image(
        image_path=args.out_png,
        out_docx_path=args.out_docx,
        title=args.title,
    )

    print(f"[OK] Created stego PNG: {args.out_png}")
    print(f"[OK] Created DOCX container: {args.out_docx}")
    print(f"[i] Cover capacity ~ {lsb_capacity_bytes(Image.open(args.cover_png).convert('RGB'))} bytes (for payload+len).")
    print(f"[i] Embedded payload bytes (packed): {len(payload)}")


def cmd_extract(args: argparse.Namespace) -> None:
    png_bytes = extract_first_png_from_docx(args.in_docx)
    stego_png_path = args.tmp_png

    with open(stego_png_path, "wb") as f:
        f.write(png_bytes)

    payload = extract_lsb_png(stego_png_path)
    packet = CryptoPacket.unpack(payload)
    plaintext = decrypt(packet, args.password)

    if args.outfile:
        with open(args.outfile, "wb") as f:
            f.write(plaintext)
        print(f"[OK] Extracted plaintext to: {args.outfile}")
    else:
        # Пытаемся вывести как UTF‑8; если не текст — печатаем «как есть»
        try:
            print(plaintext.decode("utf-8"))
        except UnicodeDecodeError:
            print(plaintext)


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="DOCX stego prototype (AES-GCM + PNG LSB).")
    sub = p.add_subparsers(dest="cmd", required=True)

    pe = sub.add_parser("embed", help="Encrypt and embed into PNG, then create DOCX with that image.")
    pe.add_argument("--cover-png", required=True, help="Path to cover PNG (lossless).")
    group = pe.add_mutually_exclusive_group(required=True)
    group.add_argument("--message", help="UTF-8 text to embed.")
    group.add_argument("--infile", help="Binary file to embed.")
    pe.add_argument("--password", required=True, help="Password for AES-GCM key derivation (scrypt).")
    pe.add_argument("--out-png", default="stego.png", help="Output stego PNG path.")
    pe.add_argument("--out-docx", default="stego_container.docx", help="Output DOCX path.")
    pe.add_argument("--title", default="Stego Container", help="DOCX title.")

    px = sub.add_parser("extract", help="Extract PNG from DOCX, decode LSB, decrypt.")
    px.add_argument("--in-docx", required=True, help="DOCX container path.")
    px.add_argument("--password", required=True, help="Password used during embedding.")
    px.add_argument("--outfile", help="Write extracted plaintext to file (binary). If omitted, print.")
    px.add_argument("--tmp-png", default="_extracted.png", help="Temp path to store extracted PNG.")

    return p


def main():
    parser = build_parser()
    args = parser.parse_args()

    try:
        if args.cmd == "embed":
            cmd_embed(args)
        elif args.cmd == "extract":
            cmd_extract(args)
        else:
            parser.error("Unknown command")
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
